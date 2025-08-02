import json
from datetime import datetime
import pandas as pd
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os

class PDFExcelGen:
    """
    A class for generating PDF, DOCX, and Excel files from question data.
    """
    
    def __init__(self, output_directory="./exports"):
        """
        Initialize the generator with output directory.
        
        Args:
            output_directory: Directory where files will be saved
        """
        self.output_directory = output_directory
        os.makedirs(output_directory, exist_ok=True)
        
    def _prepare_question_data(self, questions_data):
        """
        Prepare question data for export by flattening and formatting.
        
        Args:
            questions_data: Dictionary containing questions list
            
        Returns:
            List of formatted question dictionaries
        """
        questions = questions_data.get('questions', [])
        formatted_questions = []
        
        for question in questions:
            formatted_question = {
                'ID': question.get('question_id', ''),
                'Question': question.get('question', ''),
                'Type': question.get('type', ''),
                'Difficulty': question.get('difficulty', ''),
                'Language': question.get('language', 'English'),
                'Solution': question.get('solution', ''),
                'Tags': ', '.join(question.get('tags', [])),
                'Image Required': 'Yes' if question.get('image_required', False) else 'No'
            }
            formatted_questions.append(formatted_question)
            
        return formatted_questions
    
    def generate_excel(self, questions_data, filename=None):
        """
        Generate an Excel file from questions data.
        
        Args:
            questions_data: Dictionary containing questions list
            filename: Optional custom filename
            
        Returns:
            String path to the generated file
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"questions_export_{timestamp}.xlsx"
        
        filepath = os.path.join(self.output_directory, filename)
        
        # Prepare data for DataFrame
        formatted_questions = self._prepare_question_data(questions_data)
        
        if not formatted_questions:
            # Create empty DataFrame with headers
            df = pd.DataFrame(columns=['ID', 'Question', 'Type', 'Difficulty', 'Language', 'Solution', 'Tags', 'Image Required'])
        else:
            df = pd.DataFrame(formatted_questions)
        
        # Create Excel writer with formatting
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Questions', index=False)
            
            # Get the workbook and worksheet
            workbook = writer.book
            worksheet = writer.sheets['Questions']
            
            # Apply formatting
            from openpyxl.styles import Font, PatternFill, Alignment
            from openpyxl.utils import get_column_letter
            
            # Header formatting
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            
            for col_num in range(1, len(df.columns) + 1):
                cell = worksheet.cell(row=1, column=col_num)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
            
            # Auto-adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        return filepath
    
    def generate_pdf(self, questions_data, filename=None):
        """
        Generate a PDF file from questions data.
        
        Args:
            questions_data: Dictionary containing questions list
            filename: Optional custom filename
            
        Returns:
            String path to the generated file
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"questions_export_{timestamp}.pdf"
        
        filepath = os.path.join(self.output_directory, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph("Question Bank Export", title_style))
        story.append(Spacer(1, 12))
        
        # Metadata
        questions = questions_data.get('questions', [])
        metadata_style = ParagraphStyle(
            'Metadata',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=20
        )
        story.append(Paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", metadata_style))
        story.append(Paragraph(f"Total Questions: {len(questions)}", metadata_style))
        story.append(Spacer(1, 20))
        
        # Questions
        question_style = ParagraphStyle(
            'QuestionStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leftIndent=20
        )
        
        answer_style = ParagraphStyle(
            'AnswerStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            leftIndent=30,
            textColor=colors.blue
        )
        
        meta_style = ParagraphStyle(
            'MetaStyle',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=12,
            leftIndent=20,
            textColor=colors.grey
        )
        
        for i, question in enumerate(questions, 1):
            # Question number and text
            story.append(Paragraph(f"<b>Q{i}. {question.get('question', '')}</b>", question_style))
            
            # Solution
            if question.get('solution'):
                story.append(Paragraph(f"<b>Answer:</b> {question.get('solution', '')}", answer_style))
            
            # Metadata
            meta_info = []
            meta_info.append(f"Type: {question.get('type', 'N/A')}")
            meta_info.append(f"Difficulty: {question.get('difficulty', 'N/A')}")
            meta_info.append(f"Language: {question.get('language', 'English')}")
            if question.get('tags'):
                meta_info.append(f"Tags: {', '.join(question.get('tags', []))}")
            
            story.append(Paragraph(" | ".join(meta_info), meta_style))
            story.append(Spacer(1, 12))
            
            # Page break after every 5 questions
            if i % 5 == 0 and i < len(questions):
                story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        return filepath
    
    def generate_docx(self, questions_data, filename=None):
        """
        Generate a DOCX file from questions data.
        
        Args:
            questions_data: Dictionary containing questions list
            filename: Optional custom filename
            
        Returns:
            String path to the generated file
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"questions_export_{timestamp}.docx"
        
        filepath = os.path.join(self.output_directory, filename)
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('Question Bank Export', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        questions = questions_data.get('questions', [])
        doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Questions: {len(questions)}")
        doc.add_paragraph("")  # Empty line
        
        # Questions
        for i, question in enumerate(questions, 1):
            # Question
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(f"Q{i}. {question.get('question', '')}")
            question_run.bold = True
            
            # Solution
            if question.get('solution'):
                answer_para = doc.add_paragraph()
                answer_run = answer_para.add_run("Answer: ")
                answer_run.bold = True
                answer_para.add_run(question.get('solution', ''))
            
            # Metadata table
            metadata_table = doc.add_table(rows=1, cols=2)
            metadata_table.style = 'Table Grid'
            
            # Add metadata rows
            metadata_items = [
                ('Type', question.get('type', 'N/A')),
                ('Difficulty', question.get('difficulty', 'N/A')),
                ('Language', question.get('language', 'English')),
                ('Image Required', 'Yes' if question.get('image_required', False) else 'No'),
                ('Tags', ', '.join(question.get('tags', [])) if question.get('tags') else 'None')
            ]
            
            for key, value in metadata_items:
                row_cells = metadata_table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
            
            # Add spacing
            doc.add_paragraph("")
            
            # Page break after every 3 questions
            if i % 3 == 0 and i < len(questions):
                doc.add_page_break()
        
        # Save document
        doc.save(filepath)
        return filepath
    
    def generate_all_formats(self, questions_data, base_filename=None):
        """
        Generate files in all formats (PDF, DOCX, Excel).
        
        Args:
            questions_data: Dictionary containing questions list
            base_filename: Base filename (without extension)
            
        Returns:
            Dictionary containing paths to all generated files
        """
        if base_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"questions_export_{timestamp}"
        
        results = {}
        
        try:
            results['excel'] = self.generate_excel(questions_data, f"{base_filename}.xlsx")
        except Exception as e:
            results['excel_error'] = str(e)
        
        try:
            results['pdf'] = self.generate_pdf(questions_data, f"{base_filename}.pdf")
        except Exception as e:
            results['pdf_error'] = str(e)
        
        try:
            results['docx'] = self.generate_docx(questions_data, f"{base_filename}.docx")
        except Exception as e:
            results['docx_error'] = str(e)
        
        return results
    
    def get_file_stats(self, questions_data):
        """
        Get statistics about the questions data.
        
        Args:
            questions_data: Dictionary containing questions list
            
        Returns:
            Dictionary containing statistics
        """
        questions = questions_data.get('questions', [])
        
        stats = {
            'total_questions': len(questions),
            'difficulty_distribution': {},
            'type_distribution': {},
            'language_distribution': {},
            'questions_with_images': 0,
            'total_tags': set()
        }
        
        for question in questions:
            # Difficulty distribution
            difficulty = question.get('difficulty', 'Unknown')
            stats['difficulty_distribution'][difficulty] = stats['difficulty_distribution'].get(difficulty, 0) + 1
            
            # Type distribution
            question_type = question.get('type', 'Unknown')
            stats['type_distribution'][question_type] = stats['type_distribution'].get(question_type, 0) + 1
            
            # Language distribution
            language = question.get('language', 'English')
            stats['language_distribution'][language] = stats['language_distribution'].get(language, 0) + 1
            
            # Image requirements
            if question.get('image_required', False):
                stats['questions_with_images'] += 1
            
            # Tags
            for tag in question.get('tags', []):
                stats['total_tags'].add(tag)
        
        stats['total_unique_tags'] = len(stats['total_tags'])
        stats['total_tags'] = list(stats['total_tags'])
        
        return stats

# Example usage
if __name__ == "__main__":
    # Sample data
    sample_data = {
        "questions": [
            {
                "question_id": 1,
                "question": "What is the capital of France?",
                "type": "MCQ",
                "difficulty": "Easy",
                "language": "English",
                "solution": "Paris",
                "tags": ["geography", "europe"],
                "image_required": False
            },
            {
                "question_id": 2,
                "question": "Explain the process of photosynthesis.",
                "type": "Long Answer",
                "difficulty": "Medium",
                "language": "English",
                "solution": "Photosynthesis is the process by which plants convert light energy into chemical energy...",
                "tags": ["biology", "plants"],
                "image_required": True
            }
        ]
    }
    
    # Create generator
    generator = PDFExcelGen()
    
    # Generate all formats
    results = generator.generate_all_formats(sample_data)
    print("Generated files:", results)
    
    # Get statistics
    stats = generator.get_file_stats(sample_data)
    print("Statistics:", stats)
